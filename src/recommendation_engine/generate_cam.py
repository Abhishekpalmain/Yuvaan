import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class RecommendationEngine:
    """
    Pillar 3: Recommendation Engine
    Synthesizes inputs from all pillars and generates explainable decisions
    """

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.scoring_config = config.get('scoring', {})
        self.decisions_config = config.get('decisions', {})

    def evaluate(self,
                 company_name: str,
                 gst_results: Optional[Dict[str, Any]] = None,
                 circular_results: Optional[Dict[str, Any]] = None,
                 cibil_rank: Optional[int] = None,
                 research_findings: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Evaluate loan application based on all pillar inputs

        Args:
            company_name: Company name
            gst_results: Output from GSTR reconciliation
            circular_results: Output from circular trading detection
            cibil_rank: CIBIL commercial rank (1-10)
            research_findings: Output from research agent

        Returns:
            Evaluation results dictionary
        """
        logger.info(f"Starting evaluation for {company_name}")

        # Start with base score
        base_score = self.scoring_config.get('base_score', 80)
        final_score = base_score
        explanations = []

        evaluation = {
            "company_name": company_name,
            "timestamp": datetime.now().isoformat(),
            "base_score": base_score,
            "adjustments": [],
            "final_score": final_score,
            "recommendation": "",
            "credit_limit": "INR 0",
            "interest_rate": "N/A",
            "risk_factors": [],
            "positive_factors": []
        }

        # Pillar 1: GST Reconciliation Check
        if gst_results and gst_results.get('summary', {}).get('has_discrepancy', False):
            deduction = self.scoring_config.get('deductions', {}).get('gst_discrepancy', -20)
            final_score += deduction
            discrepancies_count = len(gst_results['discrepancies'])
            msg = f"DEDUCTION ({deduction:+d}): GST 3B claims exceed 2A supplier filings. {discrepancies_count} discrepancy(ies) found. Potential tax evasion risk."
            explanations.append(msg)
            evaluation['adjustments'].append({
                "type": "gst_discrepancy",
                "value": deduction,
                "details": gst_results['discrepancies']
            })
            evaluation['risk_factors'].append("GST ITC inflation detected")
            logger.warning(f"Score adjustment: {msg}")

        # Pillar 1: Circular Trading Check
        if circular_results and circular_results.get('cycles_detected', False):
            deduction = self.scoring_config.get('deductions', {}).get('circular_trading', -50)
            final_score += deduction
            cycle_count = len(circular_results.get('cycles', []))
            msg = f"DEDUCTION ({deduction:+d}): Circular trading loop(s) detected in Bank Statements. High fraud risk. {cycle_count} loop(s) identified."
            explanations.append(msg)
            evaluation['adjustments'].append({
                "type": "circular_trading",
                "value": deduction,
                "cycles": circular_results.get('cycles', [])
            })
            evaluation['risk_factors'].append("Circular trading detected")
            logger.warning(f"Score adjustment: {msg}")

        # CIBIL Rank Check
        if cibil_rank is not None:
            if cibil_rank > 7:
                deduction = self.scoring_config.get('deductions', {}).get('poor_cibil', -10)
                final_score += deduction
                msg = f"DEDUCTION ({deduction:+d}): Poor CIBIL Commercial Rank (Rank {cibil_rank}/10)."
                explanations.append(msg)
                evaluation['risk_factors'].append("Poor credit history")
                logger.info(f"Score adjustment: {msg}")
            elif cibil_rank <= 3:
                addition = self.scoring_config.get('deductions', {}).get('excellent_cibil', +10)
                final_score += addition
                msg = f"ADDITION ({addition:+d}): Excellent CIBIL Commercial Rank (Rank {cibil_rank}/10)."
                explanations.append(msg)
                evaluation['positive_factors'].append("Excellent credit history")
                logger.info(f"Score adjustment: {msg}")

        # Pillar 2: Research Findings - Litigation Risk
        if research_findings:
            litigation_risk = research_findings.get('overall_assessment', {}).get('litigation_risk', 'Low')
            sector_risk = research_findings.get('overall_assessment', {}).get('sector_risk', 'Medium')

            if litigation_risk in ['High', 'Medium-High']:
                # Auto-deduct based on research (configurable)
                research_deduction = -15  # Auto penalty for high litigation
                final_score += research_deduction
                msg = f"DEDUCTION ({research_deduction:+d}): {litigation_risk} litigation risk identified from e-Courts/MCA research."
                explanations.append(msg)
                evaluation['adjustments'].append({
                    "type": "litigation_risk",
                    "value": research_deduction,
                    "details": f"Litigation risk: {litigation_risk}"
                })
                evaluation['risk_factors'].append(f"High litigation risk ({litigation_risk})")
                logger.warning(msg)

            evaluation['research_summary'] = {
                "litigation_risk": litigation_risk,
                "sector_risk": sector_risk,
                "mca_status": research_findings.get('mca_data', {}).get('status', 'N/A'),
                "pending_cases": research_findings.get('litigation_data', {}).get('total_cases', 0)
            }

        # Clamp score to 0-100
        final_score = max(0, min(100, final_score))
        evaluation['final_score'] = final_score

        # Determine recommendation
        recommendation, limit, rate = self._get_recommendation(final_score)
        evaluation['recommendation'] = recommendation
        evaluation['credit_limit'] = limit
        evaluation['interest_rate'] = rate

        # Generate final summary
        evaluation['summary'] = self._generate_summary(evaluation, explanations)

        logger.info(f"Evaluation complete: Score={final_score}, Recommendation={recommendation}")

        return evaluation

    def _get_recommendation(self, score: int) -> tuple:
        """Determine recommendation based on score"""
        approve_min = self.decisions_config.get('approve_min_score', 70)
        conditional_min = self.decisions_config.get('conditional_min_score', 50)

        if score >= approve_min:
            return (
                "APPROVE",
                self.decisions_config.get('approve_min_limit_formatted', "INR 5,00,00,000"),
                self.decisions_config.get('approve_rate', "10.5% p.a.")
            )
        elif score >= conditional_min:
            return (
                "APPROVE WITH CONDITIONS",
                self.decisions_config.get('conditional_max_limit_formatted', "INR 2,00,00,000"),
                self.decisions_config.get('conditional_rate', "12.0% p.a. (Risk Premium Applied)")
            )
        else:
            return ("REJECT", "INR 0", "N/A")

    def _generate_summary(self, evaluation: Dict[str, Any], explanations: list) -> str:
        """Generate executive summary"""
        risk_count = len(evaluation['risk_factors'])
        pos_count = len(evaluation['positive_factors'])

        if evaluation['recommendation'] == "APPROVE":
            summary = f"APPROVAL RECOMMENDED. AI Credit Score: {evaluation['final_score']}/100. "
            if risk_count == 0:
                summary += "No significant risk factors identified."
            else:
                summary += f"Minor risk factors noted but overall profile is strong."
        elif evaluation['recommendation'] == "APPROVE WITH CONDITIONS":
            summary = f"CONDITIONAL APPROVAL. AI Credit Score: {evaluation['final_score']}/100. "
            summary += f"Caution advised due to {risk_count} risk factor(s). "
            summary += "Recommendation: approve with reduced limit and risk premium."
        else:
            summary = f"REJECTION RECOMMENDED. AI Credit Score: {evaluation['final_score']}/100. "
            summary += f"Critical risk factors identified ({risk_count} total). "
            summary += "Unable to recommend credit facility at this time."

        return summary

    def generate_cam_document(self, evaluation: Dict[str, Any], output_path: str = None) -> str:
        """
        Generate Word document (CAM) from evaluation results

        Args:
            evaluation: Evaluation results dictionary
            output_path: Optional output directory path

        Returns:
            Path to generated document
        """
        if output_path is None:
            output_path = self.config['paths']['output_dir']

        doc = Document()
        company_name = evaluation['company_name']

        # Title
        title = doc.add_heading('Credit Appraisal Memo (CAM)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(f"Applicant: {company_name}")
        doc.add_paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Final AI Credit Score: {evaluation['final_score']}/100")

        # Recommendation (highlighted)
        para = doc.add_paragraph()
        run = para.add_run(f"Recommendation: {evaluation['recommendation']}")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0) if evaluation['recommendation'] == "REJECT" else RGBColor(0, 128, 0)
        para.style = 'Intense Quote'

        doc.add_paragraph(f"Suggested Credit Limit: {evaluation['credit_limit']}")
        doc.add_paragraph(f"Suggested Interest Rate: {evaluation['interest_rate']}")
        doc.add_paragraph(f"Summary: {evaluation['summary']}")

        # AI Decision Explanation
        doc.add_heading('2. AI Decision Explanation (Explainable AI)', level=1)
        intro = doc.add_paragraph("This section provides complete transparency into the AI decision-making process. Every score adjustment is documented below with explicit reasoning:")
        intro.paragraph_format.space_after = Pt(12)

        for adj in evaluation.get('adjustments', []):
            if adj['type'] == 'gst_discrepancy':
                for detail in adj.get('details', []):
                    msg = detail.get('message', 'GST discrepancy detected')
                    doc.add_paragraph(msg, style='List Bullet')
            elif adj['type'] == 'circular_trading':
                for cycle in adj.get('cycles', []):
                    path = " -> ".join(cycle['path'])
                    msg = f"Circular trading loop detected: {path} -> {cycle['path'][0]} ({len(cycle['transactions'])} transactions)"
                    doc.add_paragraph(msg, style='List Bullet')
            else:
                # Generic adjustment
                doc.add_paragraph(f"Score adjustment: {adj['value']:+d}", style='List Bullet')

        # The Five Cs of Credit
        doc.add_heading('3. The Five Cs of Credit', level=1)

        # Character
        p = doc.add_paragraph()
        p.add_run("Character: ").bold = True
        char_text = "Auto-filled from research\n"
        if evaluation.get('research_summary'):
            char_text += f"- MCA Status: {evaluation['research_summary'].get('mca_status', 'N/A')}\n"
            char_text += f"- Pending Litigation: {evaluation['research_summary'].get('pending_cases', 0)} cases\n"
            char_text += f"- Litigation Risk: {evaluation['research_summary'].get('litigation_risk', 'N/A')}"
        p.add_run(char_text)

        # Capacity
        p = doc.add_paragraph()
        p.add_run("Capacity: ").bold = True
        cap_text = "Auto-filled from GST & Bank Statement analysis\n"
        if evaluation.get('risk_factors'):
            if any('GST' in factor for factor in evaluation['risk_factors']):
                cap_text += "- WARNING: GST ITC inflation detected (see explanations above)\n"
            if any('circular' in factor.lower() for factor in evaluation['risk_factors']):
                cap_text += "- WARNING: Circular trading patterns detected in bank transactions\n"
        if not any(c in str(evaluation.get('risk_factors', [])) for c in ['GST', 'circular']):
            cap_text += "- Cash flow analysis shows adequate service capacity"
        p.add_run(cap_text)

        # Capital
        p = doc.add_paragraph()
        p.add_run("Capital: ").bold = True
        p.add_run("Auto-filled from Annual Reports (OCR-extracted)")

        # Collateral
        p = doc.add_paragraph()
        p.add_run("Collateral: ").bold = True
        p.add_run("Pending physical valuation and title verification")

        # Conditions
        p = doc.add_paragraph()
        p.add_run("Conditions: ").bold = True
        p.add_run(f"Standard corporate lending rates. Current recommendation based on {evaluation['final_score']}/100 risk score.")

        # Risk Factors Summary
        if evaluation.get('risk_factors'):
            doc.add_heading('4. Identified Risk Factors', level=1)
            for factor in evaluation['risk_factors']:
                doc.add_paragraph(factor, style='List Bullet')

        # Positive Factors
        if evaluation.get('positive_factors'):
            doc.add_heading('5. Positive Factors', level=1)
            for factor in evaluation['positive_factors']:
                doc.add_paragraph(factor, style='List Bullet')

        # Footer
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Generated by Intelli-Credit Decisioning Engine | ").font.size = Pt(9)
        footer.add_run("This is an AI-assisted recommendation. Human approval required.").font.size = Pt(9)

        # Save document
        safe_company_name = company_name.replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{output_path}/CAM_{safe_company_name}.docx"
        doc.save(filename)

        logger.info(f"CAM document generated: {filename}")
        print(f"[OK] CAM document saved to: {filename}")

        return filename


def generate_cam(company_name: str,
                 base_score: int,
                 gst_risk: bool,
                 circular_risk: bool,
                 cibil_rank: int,
                 config_path: str = "config.yaml") -> str:
    """
    Main entry point for CAM generation (backward compatibility)

    Args:
        company_name: Company name
        base_score: Starting credit score
        gst_risk: Whether GST discrepancies exist
        circular_risk: Whether circular trading detected
        cibil_rank: CIBIL rank 1-10
        config_path: Path to config file

    Returns:
        Path to generated document
    """
    import yaml
    config = yaml.safe_load(open(config_path))

    engine = RecommendationEngine(config)
    evaluation = engine.evaluate(
        company_name=company_name,
        gst_results={"summary": {"has_discrepancy": gst_risk}, "discrepancies": []} if gst_risk else None,
        circular_results={"cycles_detected": circular_risk, "cycles": []} if circular_risk else None,
        cibil_rank=cibil_rank
    )

    return engine.generate_cam_document(evaluation)


if __name__ == "__main__":
    import yaml
    config = yaml.safe_load(open("config.yaml"))

    engine = RecommendationEngine(config)

    # For demo, use hardcoded values
    evaluation = engine.evaluate(
        company_name=config['company']['name'],
        gst_results={"summary": {"has_discrepancy": True}, "discrepancies": [{"message": "Mock discrepancy"}]},
        circular_results={"cycles_detected": True, "cycles": []},
        cibil_rank=4
    )

    engine.generate_cam_document(evaluation)

    print("\n=== EVALUATION SUMMARY ===")
    print(f"Final Score: {evaluation['final_score']}/100")
    print(f"Recommendation: {evaluation['recommendation']}")
    print(f"Credit Limit: {evaluation['credit_limit']}")
    print(f"Interest Rate: {evaluation['interest_rate']}")

